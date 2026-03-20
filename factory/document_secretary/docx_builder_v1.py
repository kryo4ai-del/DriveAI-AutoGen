"""Professional .docx document builder with consistent styling."""

from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


class DocxBuilder:
    """Builds professional .docx documents with consistent styling."""

    def __init__(self, title: str, subtitle: str = ""):
        self.doc = Document()
        self._setup_styles()
        self._add_title_page(title, subtitle)

    def _setup_styles(self):
        """Configure document-wide styles: A4, professional fonts."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15

        for level, size, color in [
            (1, 18, "1a1a2e"),
            (2, 14, "16213e"),
            (3, 12, "0f3460"),
        ]:
            heading_style = self.doc.styles[f"Heading {level}"]
            heading_style.font.name = "Calibri"
            heading_style.font.size = Pt(size)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor.from_string(color)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)

    def _add_title_page(self, title: str, subtitle: str):
        """Add a clean title page."""
        for _ in range(6):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Stand: {date.today().strftime('%d. %B %Y')}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("DriveAI Swarm Factory")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        self.doc.add_page_break()

    def add_heading(self, text: str, level: int = 1):
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    def add_key_value(self, key: str, value: str):
        """Add a key-value line like 'Idee: EchoMatch'."""
        p = self.doc.add_paragraph()
        run_key = p.add_run(f"{key}: ")
        run_key.bold = True
        p.add_run(value)

    def add_table(self, headers: list[str], rows: list[list[str]]):
        """Add a formatted table."""
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                row.cells[i].text = str(cell_text)
                for paragraph in row.cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

        self.doc.add_paragraph()

    def add_traffic_light(self, field: str, status: str, details: str = ""):
        """Add a traffic light risk indicator."""
        p = self.doc.add_paragraph()
        run = p.add_run(f"{status} {field}")
        run.bold = True
        if details:
            p.add_run(f" — {details}")

    def add_page_break(self):
        self.doc.add_page_break()

    def add_section_separator(self):
        """Add a subtle line separator."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("—  —  —")
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    def save(self, filepath: str) -> str:
        self.doc.save(filepath)
        return filepath
